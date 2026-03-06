import streamlit as st
import datetime
import requests
import pandas as pd
import io
import pytz
import time
from docx import Document

# 1. Page Configuration
st.set_page_config(page_title="NGX Market Dashboard", layout="wide")

# --- TIMEZONE LOGIC ---
def get_wat_time():
    """Returns the current time in West Africa Time (Lagos)"""
    utc_now = datetime.datetime.now(datetime.timezone.utc)
    wat_tz = pytz.timezone('Africa/Lagos')
    return utc_now.astimezone(wat_tz)

# 2. DATA FETCHING 
@st.cache_data(ttl=60)
def get_ngx_api_data(endpoint):
    url = f"https://doclib.ngxgroup.com/REST/api/mrkstat/{endpoint}"
    headers = {"User-Agent": "Mozilla/5.0", "Referer": "https://ngxgroup.com/"}
    try:
        response = requests.get(url, headers=headers, timeout=10)
        data = response.json()
        output = []
        for item in data[:7]:
            last_close = float(item.get('LAST_CLOSE', 0))
            price_change = float(item.get('PERCENTAGE_CHANGE', 0))
            todays_close = float(item.get('TODAYS_CLOSE', 0))
            pc_val = (price_change / last_close * 100) if last_close != 0 else 0
            
            output.append({
                "Symbol": item.get('SYMBOL', 'N/A'),
                "Price": todays_close,
                "Change (N)": price_change,
                "% Change": round(pc_val, 2)
            })
        return pd.DataFrame(output)
    except:
        return pd.DataFrame()

# 3. HEADER & LIVE CLOCK PLACEHOLDER
st.title("NGX Daily Equity Summary")

clock_placeholder = st.empty()

# 4. MAIN DASHBOARD CONTENT
col1, col2 = st.columns(2)


gainers_df = get_ngx_api_data("topsymbols")
losers_df = get_ngx_api_data("bottomsymbols")

with col1:
    st.success("**Top 7 Advancers**")
    if not gainers_df.empty:
        st.dataframe(gainers_df, use_container_width=True, hide_index=True)
    else:
        st.warning("No gainers data available currently.")

with col2:
    st.error("**Top 7 Decliners**")
    if not losers_df.empty:
        st.dataframe(losers_df, use_container_width=True, hide_index=True)
    else:
        st.warning("No losers data available currently.")

# 5. DOWNLOAD SECTION
st.divider()
st.subheader("Export Reports")

btn_col1, btn_col2 = st.columns(2)

# --- EXCEL DOWNLOAD LOGIC ---
if not gainers_df.empty or not losers_df.empty: 
    excel_bio = io.BytesIO()
    with pd.ExcelWriter(excel_bio, engine='openpyxl') as writer:
        if not gainers_df.empty:
            gainers_df.to_excel(writer, sheet_name='Top Gainers', index=False)
        if not losers_df.empty:
            losers_df.to_excel(writer, sheet_name='Top Losers', index=False)
    
    excel_bio.seek(0)
    
    with btn_col1:
        st.download_button(
            label="Download Excel Report",
            data=excel_bio.getvalue(),
            file_name=f"NGX_Report_{get_wat_time().strftime('%Y-%m-%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
# --- WORD DOWNLOAD LOGIC ---

from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- STYLING HELPERS ---
def set_cell_background(cell, fill_color):
    """Sets cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_padding(cell, top=100, bottom=100):
    """Adds vertical padding to the cell (dxa units)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'bottom']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(top if margin == 'top' else bottom))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_table_font(cell, font_name, size, is_bold=False):
    """Sets font, size, and forces text to WHITE inside colored cells."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            run.bold = is_bold
            run.font.color.rgb = RGBColor(255, 255, 255) # White text for contrast
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)

# --- WORD GENERATION ---
if not gainers_df.empty or not losers_df.empty:
    doc = Document()
    current_wat = get_wat_time()
    
    # 1. NGX Market Summary: Calibri 12, Black, No Blue Underline
    # We use a standard paragraph instead of add_heading to avoid default blue styles
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('NGX Market Summary')
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14) # Slightly larger for a main title
    title_run.font.color.rgb = RGBColor(0, 0, 0) # Black text

    date_para = doc.add_paragraph(f"Report Date: {current_wat.strftime('%d %b %Y at %I:%M:%S %p WAT')}")
    date_run = date_para.runs[0]
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0, 0, 0) # Black text

    # 2. TOP GAINERS SECTION
    if not gainers_df.empty:
        # Section Heading: Calibri 12, White (assuming user wants this contrast)
        g_head = doc.add_paragraph()
        g_run = g_head.add_run('Top Gainers')
        g_run.bold = True
        g_run.font.name = 'Calibri'
        g_run.font.size = Pt(12)
        g_run.font.color.rgb = RGBColor(255, 255, 255) # White heading text

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid' # Ensures Black Borders
        
        # Header Row: Calibri Bold 11, Green
        cols = ["Gainers", "Close Price", "% Change", "Naira Change"]
        for i, text in enumerate(cols):
            cell = table.rows[0].cells[i]
            cell.text = text
            set_cell_background(cell, "8EB77D") # Muted Green
            set_cell_padding(cell, top=140, bottom=140) # Extra padding
            apply_table_font(cell, "Calibri", 11, is_bold=True)

        # Body Rows: Aptos 11, Green
        for _, row in gainers_df.iterrows():
            row_cells = table.add_row().cells
            vals = [row['Symbol'], f"{row['Price']:.2f}", f"{row['% Change']:.2f}", f"{row['Change (N)']:.2f}"]
            for i, val in enumerate(vals):
                row_cells[i].text = str(val)
                set_cell_background(row_cells[i], "8EB77D")
                set_cell_padding(row_cells[i], top=100, bottom=100)
                apply_table_font(row_cells[i], "Aptos", 11)

    doc.add_paragraph("\n") 

    # 3. TOP LOSERS SECTION
    if not losers_df.empty:
        l_head = doc.add_paragraph()
        l_run = l_head.add_run('Top Losers')
        l_run.bold = True
        l_run.font.name = 'Calibri'
        l_run.font.size = Pt(12)
        l_run.font.color.rgb = RGBColor(255, 255, 255) # White heading text
        
        table_l = doc.add_table(rows=1, cols=4)
        table_l.style = 'Table Grid' # Ensures Black Borders
        
        cols_l = ["Losers", "Close Price", "% Change", "Naira Change"]
        for i, text in enumerate(cols_l):
            cell = table_l.rows[0].cells[i]
            cell.text = text
            set_cell_background(cell, "EB5952") # Coral Red
            set_cell_padding(cell, top=140, bottom=140)
            apply_table_font(cell, "Calibri", 11, is_bold=True)

        for _, row in losers_df.iterrows():
            row_cells = table_l.add_row().cells
            vals = [row['Symbol'], f"{row['Price']:.2f}", f"{row['% Change']:.2f}", f"{row['Change (N)']:.2f}"]
            for i, val in enumerate(vals):
                row_cells[i].text = str(val)
                set_cell_background(row_cells[i], "EB5952")
                set_cell_padding(row_cells[i], top=100, bottom=100)
                apply_table_font(row_cells[i], "Aptos", 11)

    # Save and Export
    word_bio = io.BytesIO()
    doc.save(word_bio)
    with btn_col2:
        st.download_button(
            label="Download Word Report",
            data=word_bio.getvalue(),
            file_name=f"NGX_Report_{current_wat.strftime('%Y-%m-%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# 6. LIVE CLOCK 
while True:
    now = get_wat_time()
    clock_placeholder.markdown(
        f" **{now.strftime('%d %b %Y')}** | **{now.strftime('%I:%M:%S %p')} WAT**"
    )
    time.sleep(1)
